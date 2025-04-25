import streamlit as st
from docx import Document
import io
import zipfile

# Fungsi untuk kompresi (Run Length Encoding)
def compress(data):
    compressed_data = []
    i = 0
    while i < len(data):
        char = data[i]
        count = 1
        while i + 1 < len(data) and data[i + 1] == char:
            i += 1
            count += 1
        # Simpan posisi dan jumlah pengulangan
        compressed_data.append((i, char, count))
        i += 1
    return compressed_data

# Fungsi untuk dekompresi (Run Length Decoding)
def decompress(compressed_data):
    decompressed_data = []
    for pos, char, count in compressed_data:
        decompressed_data.append(char * count)  # Rekonstruksi data
    return ''.join(decompressed_data)

# Fungsi untuk membaca file .txt
def read_txt(file):
    return file.getvalue().decode("utf-8")

# Fungsi untuk membaca file .docx
def read_docx(file):
    doc = Document(io.BytesIO(file.read()))
    return "\n".join([para.text for para in doc.paragraphs])

# Fungsi untuk menulis file .txt
def write_txt(content):
    return content.encode("utf-8")

# Fungsi untuk menulis file .docx
def write_docx(content):
    doc = Document()
    doc.add_paragraph(content)
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

# UI Streamlit
st.title('Run Length Encoding (RLE) Compression & Decompression')

# Upload file
uploaded_file = st.file_uploader("Pilih file untuk di-upload", type=["txt", "docx"])

if uploaded_file is not None:
    # Menampilkan ukuran file yang di-upload
    file_size = len(uploaded_file.getvalue())
    st.write(f"Ukuran file yang di-upload: {file_size} bytes")

    # Membaca konten file
    if uploaded_file.type == "text/plain":
        input_data = read_txt(uploaded_file)
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        input_data = read_docx(uploaded_file)
    
    # Kompresi
    st.write("### Kompresi")
    if st.button('Kompresi'):
        if input_data:
            compressed_result = compress(input_data)
            compressed_data = decompress(compressed_result)  # Dekompresi kembali untuk mendapatkan hasil asli
            compressed_file_size = len(write_txt(compressed_data) if uploaded_file.type == "text/plain" else write_docx(compressed_data))
            st.write(f"Ukuran file setelah kompresi: {compressed_file_size} bytes")

            # Menyediakan file terkompresi untuk diunduh
            if uploaded_file.type == "text/plain":
                st.download_button("Unduh file terkompresi (TXT)", write_txt(compressed_data), "compressed.txt")
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                st.download_button("Unduh file terkompresi (DOCX)", write_docx(compressed_data), "compressed.docx")
        else:
            st.warning('File kosong, tidak dapat diproses.')

    # Dekompresi
    st.write("### Dekompresi")
    if st.button('Dekompresi'):
        if input_data:
            compressed_result = compress(input_data)
            decompressed_result = decompress(compressed_result)
            decompressed_file_size = len(write_txt(decompressed_result) if uploaded_file.type == "text/plain" else write_docx(decompressed_result))
            st.write(f"Ukuran file setelah dekompresi: {decompressed_file_size} bytes")

            # Menyediakan file terdekompresi untuk diunduh
            if uploaded_file.type == "text/plain":
                st.download_button("Unduh file terdekompresi (TXT)", write_txt(decompressed_result), "decompressed.txt")
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                st.download_button("Unduh file terdekompresi (DOCX)", write_docx(decompressed_result), "decompressed.docx")
        else:
            st.warning('File kosong, tidak dapat diproses.')
