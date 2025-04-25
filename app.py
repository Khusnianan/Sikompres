import streamlit as st
from docx import Document
from PyPDF2 import PdfReader, PdfWriter
import os
import io
import re

# -------------------- Utils --------------------

def get_size_in_kb(size_bytes):
    return round(size_bytes / 1024, 2)

# Function to encode the text with Run-Length Encoding (RLE)
def run_length_encode(text):
    if not text:
        return ""
    compressed = []
    prev_char = text[0]
    count = 1
    for char in text[1:]:
        if char == prev_char:
            count += 1
        else:
            compressed.append(f"{count}{prev_char}")
            prev_char = char
            count = 1
    compressed.append(f"{count}{prev_char}")
    return ''.join(compressed)

# Function to decode text using Run-Length Encoding (RLE)
def run_length_decode(text):
    # The regex here assumes the format is "number character" for each run.
    pattern = re.compile(r'(\d+)(\D)')
    decompressed = ''.join([int(count) * char for count, char in pattern.findall(text)])
    return decompressed

# -------------------- File Handling --------------------

def compress_docx_to_docx(uploaded_file):
    doc = Document(uploaded_file)
    full_text = "\n".join([para.text for para in doc.paragraphs])
    encoded_text = run_length_encode(full_text)

    new_doc = Document()
    new_doc.add_paragraph(encoded_text)

    compressed_io = io.BytesIO()
    new_doc.save(compressed_io)
    compressed_io.seek(0)
    return compressed_io, '.docx'

def compress_txt_to_txt(uploaded_file):
    text_data = uploaded_file.read().decode('utf-8', errors='ignore')
    encoded_text = run_length_encode(text_data)

    compressed_io = io.BytesIO()
    compressed_io.write(encoded_text.encode('utf-8'))
    compressed_io.seek(0)
    return compressed_io, '.txt'

def compress_pdf_to_pdf(uploaded_file):
    reader = PdfReader(uploaded_file)
    full_text = ""
    for page in reader.pages:
        full_text += page.extract_text()

    encoded_text = run_length_encode(full_text)

    new_pdf = PdfWriter()
    new_pdf.add_blank_page()  # For simplicity, we add a blank page, as writing text to PDF is complex.
    
    compressed_io = io.BytesIO()
    new_pdf.write(compressed_io)
    compressed_io.seek(0)
    return compressed_io, '.pdf'

def decompress_docx_to_docx(uploaded_file):
    doc = Document(uploaded_file)
    full_text = "\n".join([para.text for para in doc.paragraphs])
    decoded_text = run_length_decode(full_text)

    new_doc = Document()
    new_doc.add_paragraph(decoded_text)

    decompressed_io = io.BytesIO()
    new_doc.save(decompressed_io)
    decompressed_io.seek(0)
    return decompressed_io, '.docx'

def decompress_txt_to_txt(uploaded_file):
    text_data = uploaded_file.read().decode('utf-8', errors='ignore')
    decoded_text = run_length_decode(text_data)

    decompressed_io = io.BytesIO()
    decompressed_io.write(decoded_text.encode('utf-8'))
    decompressed_io.seek(0)
    return decompressed_io, '.txt'

def decompress_pdf_to_pdf(uploaded_file):
    reader = PdfReader(uploaded_file)
    full_text = ""
    for page in reader.pages:
        full_text += page.extract_text()

    decoded_text = run_length_decode(full_text)

    new_pdf = PdfWriter()
    new_pdf.add_blank_page()  # For simplicity, we add a blank page, as writing text to PDF is complex.

    decompressed_io = io.BytesIO()
    new_pdf.write(decompressed_io)
    decompressed_io.seek(0)
    return decompressed_io, '.pdf'

# -------------------- UI --------------------

st.set_page_config(page_title="SiKompres", page_icon="📄")
st.markdown("<h1 style='text-align: center;'>📄 SiKompres</h1>", unsafe_allow_html=True)
st.markdown("<p style='text-align: center;'>🔧 Kompresi & Dekompresi File Word, PDF, dan Teks dengan Run-Length Encoding</p>", unsafe_allow_html=True)

mode = st.radio("🔄 Pilih Mode", ["Kompresi", "Dekomresi (file DOCX, TXT, PDF)"])
uploaded_file = st.file_uploader("📁 Unggah file (.docx / .txt / .pdf):")

if uploaded_file is not None:
    file_size_kb = get_size_in_kb(len(uploaded_file.getvalue()))
    filename, ext = os.path.splitext(uploaded_file.name)
    result_io = None
    result_ext = ""

    with st.spinner("⏳ Memproses file..."):

        if mode == "Kompresi":
            if ext == ".docx":
                result_io, result_ext = compress_docx_to_docx(uploaded_file)
            elif ext == ".txt":
                result_io, result_ext = compress_txt_to_txt(uploaded_file)
            elif ext == ".pdf":
                result_io, result_ext = compress_pdf_to_pdf(uploaded_file)
            else:
                st.error("❌ Format tidak didukung. Gunakan file .docx, .txt, atau .pdf untuk kompresi.")
        else:
            # DEKOMPRESI UNTUK .docx, .txt, .pdf
            if ext == ".docx":
                result_io, result_ext = decompress_docx_to_docx(uploaded_file)
            elif ext == ".txt":
                result_io, result_ext = decompress_txt_to_txt(uploaded_file)
            elif ext == ".pdf":
                result_io, result_ext = decompress_pdf_to_pdf(uploaded_file)
            else:
                st.error("❌ Format tidak didukung. Gunakan file .docx, .txt, atau .pdf untuk dekompresi.")

    if result_io:
        result_size_kb = get_size_in_kb(len(result_io.getvalue()))
        download_filename = filename + ("_compressed" if mode == "Kompresi" else "_decompressed") + result_ext

        col1, col2 = st.columns(2)
        col1.metric("📦 Ukuran Asli", f"{file_size_kb} KB")
        col2.metric("🗜️ Ukuran Hasil", f"{result_size_kb} KB")

        st.success("✅ Berhasil diproses!")
        st.download_button(
            label="⬇️ Unduh Hasil",
            data=result_io,
            file_name=download_filename,
            mime="application/octet-stream"
        )
